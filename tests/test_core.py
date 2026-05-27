import pytest
from docx import Document

from my_python_library import format_ticket, format_ticket_files, get_ticket, hello, list_ticket_files, list_ticket_numbers
from my_python_library.cli import main
from my_python_library.files import list_ticket_file_paths, open_ticket_files


def test_hello() -> None:
    assert hello("Alex") == "Hello, Alex!"


def test_ticket_numbers() -> None:
    assert list_ticket_numbers() == list(range(1, 21))


def test_get_ticket() -> None:
    ticket = get_ticket(1)
    assert ticket.number == 1
    assert "машинном обучении" in ticket.questions[0]


def test_format_ticket() -> None:
    formatted = format_ticket(20)
    assert formatted.startswith("Билет №20")
    assert "Проблемы интеллектуального права" in formatted
    assert "Ответ:" in formatted


def test_missing_ticket() -> None:
    with pytest.raises(KeyError):
        get_ticket(999)


def test_format_ticket_prefers_docx(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    folder = tmp_path / "ticket_01"
    folder.mkdir()
    document = Document()
    document.add_paragraph("Ticket from Word")
    document.add_paragraph("Question one")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "AI"
    table.cell(1, 1).text = "42"
    document.save(folder / "ticket.docx")
    (folder / "ticket.md").write_text("Ticket from Markdown", encoding="utf-8")

    monkeypatch.setattr("my_python_library.tickets.resolve_ticket_folder_path", lambda number: folder)

    formatted = format_ticket(1)
    assert "Ticket from Word" in formatted
    assert "Ticket from Markdown" not in formatted
    assert "| Term | Value |" in formatted
    assert "| AI   | 42    |" in formatted


def test_cli_prints_ticket(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["1"]) == 0
    output = capsys.readouterr().out
    assert "Билет №1" in output
    assert "Ответ:" in output
    assert "Папка файлов билета N 1" in output
    assert "ticket.md" in output
    assert any(file.filename in output for file in list_ticket_files(1) if file.filename != "ticket.md")


def test_cli_prints_ticket_without_attachment_list(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["2"]) == 0
    output = capsys.readouterr().out
    assert "Билет №2" in output
    assert "Ответ:" in output
    assert "Папка файлов билета N 2" not in output


def test_cli_lists_tickets(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["list"]) == 0
    assert "1, 2, 3" in capsys.readouterr().out


def test_cli_prints_help(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["help"]) == 0
    output = capsys.readouterr().out
    assert "Команды" in output
    assert "ist-ticket 1" in output
    assert "ist-ticket files 6" in output
    assert "ist-ticket images 6" not in output
    assert "ist-ticket open 6" in output


def test_ticket_files() -> None:
    formatted = format_ticket_files(1)
    assert "ticket_01" in formatted
    assert "ticket.md" in formatted
    for file in list_ticket_files(1):
        assert file.filename in formatted


def test_empty_ticket_files() -> None:
    formatted = format_ticket_files(2)
    assert "ticket_02" in formatted
    assert "ticket.md" in formatted


def test_cli_prints_empty_files(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["files", "2"]) == 0
    output = capsys.readouterr().out
    assert "ticket.md" in output
    assert "не добавлены" not in output


def test_cli_rejects_images_alias(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["images", "1"]) == 2
    assert "нужен номер билета" in capsys.readouterr().err


def test_open_ticket_files_tries_every_file(monkeypatch: pytest.MonkeyPatch) -> None:
    opened: list[str] = []

    def fake_open_path(path) -> str | None:
        opened.append(path.name)
        if path.name == "ticket.md":
            return "boom"
        return None

    monkeypatch.setattr("my_python_library.files.open_path", fake_open_path)
    result = open_ticket_files(1)
    expected_files = [path.name for path in list_ticket_file_paths(1)]

    for filename in expected_files:
        assert filename in opened
    assert "ticket_01" in opened
    assert result.opened == len(expected_files) - 1
    assert result.folder_opened is True
    assert result.failed == (("ticket.md", "boom"),)


def test_cli_open_reports_failures(monkeypatch: pytest.MonkeyPatch, capsys: pytest.CaptureFixture[str]) -> None:
    def fake_open_path(path) -> str | None:
        if path.name == "ticket.md":
            return "boom"
        return None

    monkeypatch.setattr("my_python_library.files.open_path", fake_open_path)
    assert main(["open", "1"]) == 0
    output = capsys.readouterr().out
    assert "Открыто файлов:" in output
    assert "Не удалось открыть:" in output
    assert "ticket.md: boom" in output


def test_cli_rejects_bad_ticket(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["999"]) == 2
    assert "не найден" in capsys.readouterr().err


def test_cli_open_rejects_bad_ticket(capsys: pytest.CaptureFixture[str]) -> None:
    assert main(["open", "999"]) == 2
    assert "не найден" in capsys.readouterr().err
